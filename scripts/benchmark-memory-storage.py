#!/usr/bin/env python3
"""
Benchmark ManageT memory + storage across all hosts running it, then
emit a Word .docx report with comparison charts.

What it measures, per host:
  - managet-agent process: RSS / VSZ / threads / uptime, taken from
    /proc/<pid>/status on Linux and from `ps` on macOS.
  - managet-agent install footprint: binary size and runtime/data
    directories (/var/run/managet, /var/log/managet if present).
  - On the dashboard host only: the Next.js production process under
    systemd (managet.service), the repo size, node_modules, .next
    build artefacts, and the SQLite DB.

Why a single Python script: keeps SSH credentials + parsing + chart
rendering + docx assembly in one place so the user can re-run after
deploys without juggling shell pipes.

Output: reports/managet-benchmark-<UTC timestamp>.docx
"""
from __future__ import annotations
import os
import re
import shutil
import subprocess
import sys
import datetime as dt
from dataclasses import dataclass, field
from pathlib import Path

# ---------------------------------------------------------------------------
# Configuration: hosts we benchmark.
#
# `dashboard_local=True` means "this machine — query directly without
# SSH and additionally collect dashboard-specific metrics". Exactly one
# entry should have it set.
# ---------------------------------------------------------------------------

REPO_ROOT = Path("/home/andrei/managet")
REPORT_DIR = REPO_ROOT / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
# SSH password for the dev hosts comes from the environment — never hardcode
# secrets in committed scripts. e.g. MANAGET_DEV_PASSWORD='…' python3 scripts/benchmark-memory-storage.py
SSH_PASSWORD = os.environ.get("MANAGET_DEV_PASSWORD")
if not SSH_PASSWORD:
    raise SystemExit("Set MANAGET_DEV_PASSWORD before running this benchmark.")
SSH_OPTS = [
    "-o", "StrictHostKeyChecking=no",
    "-o", "UserKnownHostsFile=/dev/null",
    "-o", "ConnectTimeout=8",
    "-o", "LogLevel=ERROR",
]

@dataclass
class HostSpec:
    label: str          # display name in the report
    host: str           # IP or hostname for SSH
    user: str = "andrei"
    os_kind: str = "linux"  # "linux" or "darwin"
    is_local: bool = False  # skip SSH; query this machine directly
    is_dashboard_host: bool = False  # also collect dashboard metrics

HOSTS: list[HostSpec] = [
    # The dashboard runs on the Pi — same physical box as the agent at
    # 192.168.100.82. We query locally to avoid SSHing to ourselves
    # and also collect the dashboard-only metrics.
    HostSpec(
        label="markI (Pi, Raspberry Pi 5)",
        host="127.0.0.1",
        os_kind="linux",
        is_local=True,
        is_dashboard_host=True,
    ),
    HostSpec(
        label="Mac mini (Apple Silicon)",
        host="192.168.100.95",
        os_kind="darwin",
    ),
]

# ---------------------------------------------------------------------------
# Shell helpers
# ---------------------------------------------------------------------------

def run_local(cmd: str) -> tuple[int, str]:
    p = subprocess.run(
        ["bash", "-lc", cmd],
        capture_output=True, text=True, timeout=20,
    )
    return p.returncode, (p.stdout + p.stderr).strip()


def run_remote(host: HostSpec, cmd: str) -> tuple[int, str]:
    """Run a shell command on a remote host via sshpass + ssh."""
    full = ["sshpass", "-p", SSH_PASSWORD, "ssh", *SSH_OPTS,
            f"{host.user}@{host.host}", cmd]
    p = subprocess.run(full, capture_output=True, text=True, timeout=25)
    return p.returncode, (p.stdout + p.stderr).strip()


def run_host(host: HostSpec, cmd: str) -> tuple[int, str]:
    return run_local(cmd) if host.is_local else run_remote(host, cmd)


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------

def _parse_kib(text: str, key: str) -> int | None:
    m = re.search(rf"^{re.escape(key)}:\s*(\d+)\s*kB", text, re.MULTILINE)
    return int(m.group(1)) if m else None


def _bytes_human(n: float) -> str:
    units = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while n >= 1024 and i < len(units) - 1:
        n /= 1024
        i += 1
    return f"{n:.1f} {units[i]}"


# ---------------------------------------------------------------------------
# Metric structs
# ---------------------------------------------------------------------------

@dataclass
class ProcStats:
    pid: int | None = None
    rss_kib: int | None = None        # actual RAM used
    vsz_kib: int | None = None        # virtual size (much larger, not real RAM)
    threads: int | None = None
    uptime: str | None = None         # human-readable from ps etime
    cpu_percent: float | None = None  # cumulative since process start

@dataclass
class StorageItem:
    label: str
    bytes_: int | None
    note: str = ""

@dataclass
class HostReport:
    spec: HostSpec
    os_pretty: str = ""
    kernel: str = ""
    total_ram_mib: int | None = None
    agent: ProcStats = field(default_factory=ProcStats)
    agent_storage: list[StorageItem] = field(default_factory=list)
    dashboard: ProcStats | None = None
    dashboard_storage: list[StorageItem] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Collectors — Linux
# ---------------------------------------------------------------------------

def collect_linux(host: HostSpec) -> HostReport:
    rpt = HostReport(spec=host)

    # OS info
    code, out = run_host(host, "uname -srm && cat /etc/os-release 2>/dev/null | head -5")
    if code == 0:
        lines = out.splitlines()
        if lines:
            rpt.kernel = lines[0]
        for line in lines:
            if line.startswith("PRETTY_NAME="):
                rpt.os_pretty = line.split("=", 1)[1].strip().strip('"')

    code, out = run_host(host, "grep ^MemTotal /proc/meminfo")
    if code == 0:
        m = re.search(r"(\d+)\s*kB", out)
        if m:
            rpt.total_ram_mib = int(m.group(1)) // 1024

    # Agent process
    code, pid_text = run_host(host, "pgrep -nf '^/usr/local/bin/managet-agent run' || pgrep -nx managet-agent")
    if code != 0 or not pid_text.strip().isdigit():
        rpt.errors.append("managet-agent not running")
    else:
        pid = int(pid_text.strip())
        rpt.agent.pid = pid
        code, status = run_host(host, f"cat /proc/{pid}/status")
        if code == 0:
            rpt.agent.rss_kib = _parse_kib(status, "VmRSS")
            rpt.agent.vsz_kib = _parse_kib(status, "VmSize")
            thr = re.search(r"^Threads:\s*(\d+)", status, re.MULTILINE)
            if thr:
                rpt.agent.threads = int(thr.group(1))
        code, ps = run_host(host, f"ps -o etime=,pcpu= -p {pid}")
        if code == 0 and ps.strip():
            parts = ps.strip().split()
            if len(parts) >= 2:
                rpt.agent.uptime = parts[0]
                try:
                    rpt.agent.cpu_percent = float(parts[1])
                except ValueError:
                    pass

    # Agent install / data sizes
    bin_path = "/usr/local/bin/managet-agent"
    code, sz = run_host(host, f"stat -c%s {bin_path} 2>/dev/null")
    if code == 0 and sz.strip().isdigit():
        rpt.agent_storage.append(StorageItem(
            label="Agent binary",
            bytes_=int(sz),
            note=bin_path,
        ))
    # Runtime dir (socket lives here) — usually tiny.
    code, out = run_host(host, "du -sb /var/run/managet 2>/dev/null | awk '{print $1}'")
    if code == 0 and out.strip().isdigit():
        rpt.agent_storage.append(StorageItem(
            label="Runtime dir",
            bytes_=int(out),
            note="/var/run/managet (socket)",
        ))
    code, out = run_host(host, "du -sb /var/log/managet 2>/dev/null | awk '{print $1}'")
    if code == 0 and out.strip().isdigit():
        rpt.agent_storage.append(StorageItem(
            label="Log dir",
            bytes_=int(out),
            note="/var/log/managet (rotated logs)",
        ))

    # Dashboard metrics — only on the dashboard host.
    if host.is_dashboard_host:
        rpt.dashboard = ProcStats()
        # Resolve PID via systemd to handle child re-spawns cleanly.
        code, pid_text = run_host(host, "systemctl show managet -p MainPID --value 2>/dev/null")
        pid_text = pid_text.strip()
        if pid_text.isdigit() and pid_text != "0":
            pid = int(pid_text)
            rpt.dashboard.pid = pid
            code, status = run_host(host, f"cat /proc/{pid}/status 2>/dev/null")
            if code == 0:
                rpt.dashboard.rss_kib = _parse_kib(status, "VmRSS")
                rpt.dashboard.vsz_kib = _parse_kib(status, "VmSize")
                thr = re.search(r"^Threads:\s*(\d+)", status, re.MULTILINE)
                if thr:
                    rpt.dashboard.threads = int(thr.group(1))
            code, ps = run_host(host, f"ps -o etime=,pcpu= -p {pid}")
            if code == 0 and ps.strip():
                parts = ps.strip().split()
                if len(parts) >= 2:
                    rpt.dashboard.uptime = parts[0]
                    try:
                        rpt.dashboard.cpu_percent = float(parts[1])
                    except ValueError:
                        pass
        else:
            rpt.errors.append("managet.service not running")

        # Repo size breakdown — full + the big subdirectories.
        for label, path, note in [
            ("Repo (total)", str(REPO_ROOT),
             "/home/andrei/managet"),
            ("node_modules", str(REPO_ROOT / "node_modules"),
             "third-party deps"),
            (".next build", str(REPO_ROOT / ".next"),
             "production build output"),
            ("SQLite DB", str(REPO_ROOT / "data" / "managet.db"),
             "single-file DB"),
        ]:
            code, out = run_host(host, f"du -sb {path} 2>/dev/null | awk '{{print $1}}'")
            if code == 0 and out.strip().isdigit():
                rpt.dashboard_storage.append(StorageItem(label, int(out), note))

    return rpt


# ---------------------------------------------------------------------------
# Collectors — macOS
# ---------------------------------------------------------------------------

def collect_darwin(host: HostSpec) -> HostReport:
    rpt = HostReport(spec=host)

    code, out = run_host(host, "uname -srm && sw_vers")
    if code == 0:
        lines = out.splitlines()
        if lines:
            rpt.kernel = lines[0]
        m = re.search(r"ProductName:\s*(.+)\nProductVersion:\s*(.+)", out)
        if m:
            rpt.os_pretty = f"{m.group(1).strip()} {m.group(2).strip()}"

    code, out = run_host(host, "sysctl -n hw.memsize")
    if code == 0 and out.strip().isdigit():
        rpt.total_ram_mib = int(out) // (1024 * 1024)

    # ps on macOS: -o rss,vsz are in KiB.
    code, pid_text = run_host(host, "pgrep -nx managet-agent || pgrep -nf '/managet-agent run'")
    if code != 0 or not pid_text.strip().isdigit():
        rpt.errors.append("managet-agent not running")
    else:
        pid = int(pid_text.strip())
        rpt.agent.pid = pid
        code, ps = run_host(host, f"ps -o rss=,vsz=,etime=,pcpu= -p {pid}")
        if code == 0 and ps.strip():
            parts = ps.strip().split()
            if len(parts) >= 4:
                try:
                    rpt.agent.rss_kib = int(parts[0])
                    rpt.agent.vsz_kib = int(parts[1])
                    rpt.agent.uptime = parts[2]
                    rpt.agent.cpu_percent = float(parts[3])
                except ValueError:
                    pass
        # Thread count
        code, thr = run_host(host, f"ps -M -p {pid} | tail -n +2 | wc -l")
        if code == 0 and thr.strip().isdigit():
            rpt.agent.threads = int(thr.strip())

    # Binary location is the same as Linux deployment.
    for bin_path in ("/usr/local/bin/managet-agent", "/opt/homebrew/bin/managet-agent"):
        code, sz = run_host(host, f"stat -f%z {bin_path} 2>/dev/null")
        if code == 0 and sz.strip().isdigit():
            rpt.agent_storage.append(StorageItem("Agent binary", int(sz), bin_path))
            break

    # macOS uses ~/Library/Logs by convention; agent uses /tmp or
    # /var/run depending on install. Best-effort.
    code, out = run_host(host, "du -sk /var/run/managet 2>/dev/null | awk '{print $1*1024}'")
    if code == 0 and out.strip().isdigit():
        rpt.agent_storage.append(StorageItem("Runtime dir", int(out), "/var/run/managet"))

    return rpt


# ---------------------------------------------------------------------------
# Chart rendering
# ---------------------------------------------------------------------------

def render_charts(reports: list[HostReport], out_dir: Path) -> dict[str, Path]:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    out_dir.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({
        "figure.dpi": 130,
        "savefig.dpi": 130,
        "font.size": 10,
        "axes.titleweight": "semibold",
        "axes.spines.right": False,
        "axes.spines.top": False,
    })

    paths: dict[str, Path] = {}

    # Chart 1: agent RSS comparison (MiB).
    labels = [r.spec.label for r in reports]
    rss = [(r.agent.rss_kib or 0) / 1024 for r in reports]
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    bars = ax.bar(labels, rss, color=["#a855f7", "#06b6d4"])
    ax.set_ylabel("Resident set size (MiB)")
    ax.set_title("managet-agent RAM usage by host")
    ax.set_ylim(0, max(rss + [1]) * 1.25)
    for b, v in zip(bars, rss):
        ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.1f} MiB",
                ha="center", va="bottom", fontsize=9)
    fig.tight_layout()
    p = out_dir / "agent_rss.png"
    fig.savefig(p)
    plt.close(fig)
    paths["agent_rss"] = p

    # Chart 2: agent VSZ vs RSS per host (paired bars).
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    width = 0.35
    x = list(range(len(reports)))
    rss_v = [(r.agent.rss_kib or 0) / 1024 for r in reports]
    vsz_v = [(r.agent.vsz_kib or 0) / 1024 for r in reports]
    ax.bar([i - width / 2 for i in x], rss_v, width,
           label="RSS (real RAM)", color="#a855f7")
    ax.bar([i + width / 2 for i in x], vsz_v, width,
           label="VSZ (virtual mapping)", color="#c084fc", alpha=0.7)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylabel("MiB")
    ax.set_title("agent RSS vs VSZ — RSS is what the OS actually backs with RAM")
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout()
    p = out_dir / "agent_rss_vsz.png"
    fig.savefig(p)
    plt.close(fig)
    paths["agent_rss_vsz"] = p

    # Chart 3: agent disk footprint per host.
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    bin_v = []
    other_v = []
    for r in reports:
        b = 0
        o = 0
        for s in r.agent_storage:
            if s.bytes_ is None:
                continue
            if s.label == "Agent binary":
                b += s.bytes_
            else:
                o += s.bytes_
        bin_v.append(b / (1024 * 1024))
        other_v.append(o / (1024 * 1024))
    x = list(range(len(reports)))
    ax.bar(x, bin_v, label="Binary", color="#a855f7")
    ax.bar(x, other_v, bottom=bin_v, label="Runtime + logs", color="#06b6d4")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylabel("MiB")
    ax.set_title("managet-agent on-disk footprint")
    ax.legend(frameon=False)
    for i, (b, o) in enumerate(zip(bin_v, other_v)):
        total = b + o
        if total > 0:
            ax.text(i, total, f"{total:.1f} MiB",
                    ha="center", va="bottom", fontsize=9)
    fig.tight_layout()
    p = out_dir / "agent_disk.png"
    fig.savefig(p)
    plt.close(fig)
    paths["agent_disk"] = p

    # Chart 4: dashboard repo breakdown (only one host has it).
    dash = next((r for r in reports if r.dashboard_storage), None)
    if dash:
        # Skip "Repo (total)" so the stack sums to the breakdown, not
        # double-counts.
        breakdown = [s for s in dash.dashboard_storage if s.label != "Repo (total)"]
        labels_ = [s.label for s in breakdown]
        values = [(s.bytes_ or 0) / (1024 * 1024) for s in breakdown]
        fig, ax = plt.subplots(figsize=(6.5, 3.2))
        bars = ax.bar(labels_, values,
                      color=["#a855f7", "#c084fc", "#06b6d4", "#22d3ee"])
        ax.set_ylabel("MiB")
        ax.set_title(f"Dashboard storage breakdown — {dash.spec.label}")
        for b, v in zip(bars, values):
            ax.text(b.get_x() + b.get_width() / 2, v,
                    _bytes_human(v * 1024 * 1024),
                    ha="center", va="bottom", fontsize=9)
        fig.tight_layout()
        p = out_dir / "dashboard_disk.png"
        fig.savefig(p)
        plt.close(fig)
        paths["dashboard_disk"] = p

    # Chart 5: dashboard process RAM compared to its hosting agent.
    if dash and dash.dashboard:
        fig, ax = plt.subplots(figsize=(6.5, 3.2))
        names = ["managet.service\n(dashboard)", "managet-agent\n(same host)"]
        vals = [
            (dash.dashboard.rss_kib or 0) / 1024,
            (dash.agent.rss_kib or 0) / 1024,
        ]
        bars = ax.bar(names, vals, color=["#7c3aed", "#a855f7"])
        ax.set_ylabel("Resident set size (MiB)")
        ax.set_title(f"Dashboard vs agent RAM on {dash.spec.label}")
        for b, v in zip(bars, vals):
            ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.1f} MiB",
                    ha="center", va="bottom", fontsize=9)
        fig.tight_layout()
        p = out_dir / "dashboard_vs_agent_ram.png"
        fig.savefig(p)
        plt.close(fig)
        paths["dashboard_vs_agent_ram"] = p

    return paths


# ---------------------------------------------------------------------------
# Docx assembly
# ---------------------------------------------------------------------------

def build_docx(reports: list[HostReport], charts: dict[str, Path],
               out_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Cover heading
    title = doc.add_heading("ManageT — memory & storage benchmark", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sub.add_run(
        f"Snapshot taken {dt.datetime.now(dt.timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    doc.add_paragraph(
        "This report measures live RAM and disk usage of the ManageT "
        "monitoring agent on every host where it's installed, plus the "
        "dashboard process (Next.js production server) on the host that "
        "runs it. Numbers come from /proc and du on Linux and from "
        "ps + sysctl + du on macOS. RSS is the resident set size — the "
        "memory the OS is actually backing with RAM right now — and "
        "is the figure to compare against system RAM. VSZ (virtual "
        "size) includes shared libraries and unrealized address space "
        "and is included for completeness."
    )

    # --- Summary table ---
    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 4"
    hdr = table.rows[0].cells
    for i, name in enumerate([
        "Host", "OS", "Total RAM",
        "Agent RSS", "Agent VSZ", "Agent disk"
    ]):
        hdr[i].text = name
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for rpt in reports:
        row = table.add_row().cells
        row[0].text = rpt.spec.label
        row[1].text = rpt.os_pretty or "—"
        row[2].text = (f"{rpt.total_ram_mib} MiB"
                       if rpt.total_ram_mib is not None else "—")
        row[3].text = (_bytes_human(rpt.agent.rss_kib * 1024)
                       if rpt.agent.rss_kib is not None else "—")
        row[4].text = (_bytes_human(rpt.agent.vsz_kib * 1024)
                       if rpt.agent.vsz_kib is not None else "—")
        total_disk = sum(
            (s.bytes_ or 0) for s in rpt.agent_storage
        )
        row[5].text = _bytes_human(total_disk) if total_disk else "—"

    # --- Headline charts ---
    doc.add_heading("Charts", level=1)
    doc.add_paragraph(
        "RAM usage compared across all agent installations. Bars show "
        "RSS — the actual physical memory used. The agent is written in "
        "Rust with a small tokio runtime, so a few MiB of resident "
        "memory is normal."
    )
    if "agent_rss" in charts:
        doc.add_picture(str(charts["agent_rss"]), width=Inches(6.2))
    doc.add_paragraph(
        "RSS vs VSZ — the large gap is expected. VSZ includes "
        "every memory-mapped library the agent loaded; it's not the "
        "number to budget against system RAM."
    )
    if "agent_rss_vsz" in charts:
        doc.add_picture(str(charts["agent_rss_vsz"]), width=Inches(6.2))

    doc.add_paragraph("On-disk footprint of the agent installation.")
    if "agent_disk" in charts:
        doc.add_picture(str(charts["agent_disk"]), width=Inches(6.2))

    # --- Per-host details ---
    for rpt in reports:
        doc.add_heading(rpt.spec.label, level=1)
        doc.add_paragraph(
            f"Host address: {rpt.spec.host}\n"
            f"OS: {rpt.os_pretty or '?'}\n"
            f"Kernel: {rpt.kernel or '?'}\n"
            f"Total RAM: "
            f"{rpt.total_ram_mib if rpt.total_ram_mib is not None else '?'} MiB"
        )

        if rpt.errors:
            warn = doc.add_paragraph()
            run = warn.add_run("Notes: " + "; ".join(rpt.errors))
            run.italic = True
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)

        doc.add_heading("Agent process", level=2)
        agent_lines = []
        if rpt.agent.pid is not None:
            agent_lines.append(f"PID: {rpt.agent.pid}")
        if rpt.agent.uptime:
            agent_lines.append(f"Uptime: {rpt.agent.uptime}")
        if rpt.agent.threads is not None:
            agent_lines.append(f"Threads: {rpt.agent.threads}")
        if rpt.agent.rss_kib is not None:
            agent_lines.append(
                f"RSS: {_bytes_human(rpt.agent.rss_kib * 1024)}"
                + (
                    f" — {rpt.agent.rss_kib / 1024 / rpt.total_ram_mib * 100:.2f}% of total RAM"
                    if rpt.total_ram_mib else ""
                )
            )
        if rpt.agent.vsz_kib is not None:
            agent_lines.append(f"VSZ: {_bytes_human(rpt.agent.vsz_kib * 1024)}")
        if rpt.agent.cpu_percent is not None:
            agent_lines.append(
                f"Cumulative CPU%: {rpt.agent.cpu_percent:.2f}"
            )
        doc.add_paragraph("\n".join(agent_lines) or "—")

        doc.add_heading("Agent on-disk footprint", level=2)
        if rpt.agent_storage:
            t = doc.add_table(rows=1, cols=3)
            t.style = "Light List Accent 4"
            head = t.rows[0].cells
            for i, name in enumerate(["Component", "Size", "Path / note"]):
                head[i].text = name
                for p in head[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
            for s in rpt.agent_storage:
                row = t.add_row().cells
                row[0].text = s.label
                row[1].text = _bytes_human(s.bytes_) if s.bytes_ else "—"
                row[2].text = s.note
        else:
            doc.add_paragraph("—")

        if rpt.dashboard is not None:
            doc.add_heading("Dashboard process (managet.service)", level=2)
            lines = []
            if rpt.dashboard.pid is not None:
                lines.append(f"PID: {rpt.dashboard.pid}")
            if rpt.dashboard.uptime:
                lines.append(f"Uptime: {rpt.dashboard.uptime}")
            if rpt.dashboard.threads is not None:
                lines.append(f"Threads: {rpt.dashboard.threads}")
            if rpt.dashboard.rss_kib is not None:
                lines.append(
                    f"RSS: {_bytes_human(rpt.dashboard.rss_kib * 1024)}"
                    + (
                        f" — {rpt.dashboard.rss_kib / 1024 / rpt.total_ram_mib * 100:.2f}% of total RAM"
                        if rpt.total_ram_mib else ""
                    )
                )
            if rpt.dashboard.vsz_kib is not None:
                lines.append(
                    f"VSZ: {_bytes_human(rpt.dashboard.vsz_kib * 1024)}"
                )
            if rpt.dashboard.cpu_percent is not None:
                lines.append(
                    f"Cumulative CPU%: {rpt.dashboard.cpu_percent:.2f}"
                )
            doc.add_paragraph("\n".join(lines) or "—")

            doc.add_heading("Dashboard on-disk footprint", level=2)
            if rpt.dashboard_storage:
                t = doc.add_table(rows=1, cols=3)
                t.style = "Light List Accent 4"
                head = t.rows[0].cells
                for i, name in enumerate(["Component", "Size", "Path / note"]):
                    head[i].text = name
                    for p in head[i].paragraphs:
                        for r in p.runs:
                            r.bold = True
                for s in rpt.dashboard_storage:
                    row = t.add_row().cells
                    row[0].text = s.label
                    row[1].text = _bytes_human(s.bytes_) if s.bytes_ else "—"
                    row[2].text = s.note

            doc.add_paragraph(
                "The dashboard is a Next.js 16 production server with a "
                "custom WebSocket layer for the terminal bridge. The "
                "node_modules and .next directories dominate the "
                "footprint and would be re-created on any deploy via "
                "`npm ci && npm run build`."
            )
            if "dashboard_disk" in charts:
                doc.add_picture(str(charts["dashboard_disk"]), width=Inches(6.2))
            if "dashboard_vs_agent_ram" in charts:
                doc.add_picture(str(charts["dashboard_vs_agent_ram"]),
                                width=Inches(6.2))

    # --- Methodology footer ---
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Linux hosts: agent PID via `pgrep -nf /usr/local/bin/managet-agent`; "
        "memory from /proc/<pid>/status (VmRSS, VmSize, Threads); uptime "
        "and cumulative CPU from `ps -o etime,pcpu`; total system RAM "
        "from /proc/meminfo. Disk sizes from `du -sb` (bytes)."
    )
    doc.add_paragraph(
        "macOS hosts: agent PID via `pgrep -nx managet-agent`; memory "
        "via `ps -o rss,vsz` (KiB); thread count via `ps -M`. Total RAM "
        "via `sysctl hw.memsize`. Disk sizes via `du -sk` × 1024."
    )
    doc.add_paragraph(
        "Dashboard PID is queried via `systemctl show managet -p MainPID` "
        "so the report tracks the systemd-managed process regardless of "
        "PID changes after restart."
    )
    doc.add_paragraph(
        f"Reproduce with: scripts/benchmark-memory-storage.py "
        f"(output written to reports/ as a .docx)."
    )

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    if shutil.which("sshpass") is None:
        print("error: sshpass is required", file=sys.stderr)
        return 1

    reports: list[HostReport] = []
    for host in HOSTS:
        print(f"→ collecting from {host.label} ({host.host})...")
        if host.os_kind == "linux":
            reports.append(collect_linux(host))
        elif host.os_kind == "darwin":
            reports.append(collect_darwin(host))
        else:
            print(f"  unknown os_kind: {host.os_kind}", file=sys.stderr)
            reports.append(HostReport(spec=host, errors=["unknown os_kind"]))

    charts_dir = REPORT_DIR / "_charts"
    charts = render_charts(reports, charts_dir)

    stamp = dt.datetime.now(dt.timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    out_path = REPORT_DIR / f"managet-benchmark-{stamp}.docx"
    build_docx(reports, charts, out_path)

    # Print quick recap to stdout.
    print()
    print(f"✓ wrote {out_path}")
    for rpt in reports:
        line = f"  {rpt.spec.label}: "
        if rpt.agent.rss_kib is not None:
            line += f"agent RSS={_bytes_human(rpt.agent.rss_kib * 1024)}, "
        if rpt.agent_storage:
            total = sum((s.bytes_ or 0) for s in rpt.agent_storage)
            line += f"agent disk={_bytes_human(total)}"
        if rpt.dashboard and rpt.dashboard.rss_kib is not None:
            line += f", dashboard RSS={_bytes_human(rpt.dashboard.rss_kib * 1024)}"
        print(line)
    return 0


if __name__ == "__main__":
    sys.exit(main())
